"""Genereer Word doc met gluteus medius rehab/strength programma."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT = Path(r"C:\Users\d.rietveld\Downloads\Gluteus_Medius_Programma.docx")

doc = Document()

# Marges iets smaller voor compact A4
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


def H1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return p


def H2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x5A, 0x8A)
    return p


def P(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.size = Pt(11)


def exercise_block(nummer, naam, materiaal, sets_reps, tempo, uitvoering, focus, progressie):
    H2(f"{nummer}. {naam}")
    p = doc.add_paragraph()
    p.add_run("Materiaal: ").bold = True
    p.add_run(materiaal)
    p = doc.add_paragraph()
    p.add_run("Sets/reps: ").bold = True
    p.add_run(sets_reps)
    p = doc.add_paragraph()
    p.add_run("Tempo: ").bold = True
    p.add_run(tempo)
    p = doc.add_paragraph()
    p.add_run("Uitvoering: ").bold = True
    p.add_run(uitvoering)
    p = doc.add_paragraph()
    p.add_run("Focus: ").bold = True
    p.add_run(focus)
    p = doc.add_paragraph()
    p.add_run("Progressie: ").bold = True
    p.add_run(progressie)


# ── TITEL ────────────────────────────────────────────────────────────────────
title = doc.add_heading("Gluteus Medius Programma", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Voor Dennis — opgesteld door Louis Delahaije")
run.italic = True
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Datum: 28 april 2026  |  Materiaal: weerstandsbandjes + dumbbells")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()

# ── WAAROM ───────────────────────────────────────────────────────────────────
H1("Waarom dit programma")
P(
    "Je knie-twinge bij rustige loopjes komt waarschijnlijk niet uit de knie zelf. "
    "Bij elke pas moet je gluteus medius — de spier aan de zijkant van je heup — "
    "je bekken horizontaal houden. Vuurt hij zwak of laat, dan kantelt je bekken, "
    "valt je knie naar binnen (valgus), en krijgt de patella scheef belasting. "
    "Daar voel je het."
)
P(
    "Split squats trainen vooral gluteus máximus en quadriceps. Daarom voegen we "
    "drie keer per week medius-specifiek werk toe. Tempo en vorm gaan vóór gewicht — "
    "altijd."
)
quote("\"Een sterke heup is een stille knie.\"")

# ── SCHEMA ───────────────────────────────────────────────────────────────────
H1("Weekschema")
P(
    "Doe het ACTIVATIE-blok kort vóór elke loop (5 minuten, niet vermoeiend). "
    "Doe het KRACHT-blok 2 à 3× per week op een los moment (ma/wo/vr werkt goed). "
    "Twaalf minuten per keer."
)

# ── ACTIVATIE ────────────────────────────────────────────────────────────────
H1("Activatie-blok (vóór elke loop, 5 minuten)")

exercise_block(
    nummer="A1",
    naam="Banded monster walks",
    materiaal="Weerstandsbandje rond enkels (NIET bij de knieën — dan vuurt TFL ipv medius)",
    sets_reps="2 × 10 stappen voor + 10 zij + 10 achter",
    tempo="Rustig, gecontroleerd",
    uitvoering="Lichte half-squat houding. Knieën boven enkels, niet inzakken. Kleine stappen.",
    focus="Knie blijft boven enkel — niet naar binnen vallen.",
    progressie="Zwaardere band als 2×10 makkelijk gaat.",
)

exercise_block(
    nummer="A2",
    naam="Clamshells met band",
    materiaal="Bandje rond bovenbenen, net boven de knieën",
    sets_reps="2 × 15 per kant",
    tempo="2 sec openen, 1 sec houden, 2 sec sluiten",
    uitvoering="Zijligging, knieën gebogen 90°, voeten op elkaar. Open alleen de bovenste knie. Bekken NIET mee laten roteren — alleen heup opent.",
    focus="Romp blokkeert. Het is geen torsie-oefening.",
    progressie="Zwaardere band, of houd de pauze 2 sec ipv 1 sec.",
)

exercise_block(
    nummer="A3",
    naam="Side-plank met been-lift",
    materiaal="Geen",
    sets_reps="2 × 8 per kant",
    tempo="2 sec optillen, 2 sec houden, 2 sec laten zakken",
    uitvoering="Zijplank op onderarm. Til het bovenste been gestrekt op tot iets boven heuphoogte. Bekken blijft horizontaal — niet wegzakken aan de heup-kant.",
    focus="Heup hoog en stabiel. Liever 5 schone reps dan 8 vuile.",
    progressie="Aanhouden tot 12 reps; daarna een lichte band rond enkels.",
)

# ── KRACHT ───────────────────────────────────────────────────────────────────
H1("Krachtblok (2-3× per week, ~12 min)")

exercise_block(
    nummer="K1",
    naam="Side-lying hip abduction met band of dumbbell  ★ HOOFDOEFENING",
    materiaal="Zware band rond enkels, OF dumbbell (3-5 kg) op de buitenkant van de dij vlak boven de knie",
    sets_reps="3 × 12 per been",
    tempo="2 sec omhoog — 1 sec houden — 2 sec omlaag",
    uitvoering=(
        "Zijligging, onderste been licht gebogen voor balans, bovenste been gestrekt. "
        "Til het bovenste been zijwaarts omhoog tot ongeveer 30-40°. Voet wijst recht "
        "naar voren of licht naar binnen — NIET tenen omhoog draaien (dan neemt TFL het over). "
        "Romp niet mee laten kantelen."
    ),
    focus="Voel het aan de zíjkant van de heup, niet op de zijkant van het bovenbeen of in de onderrug. Als je het verkeerd voelt: tempo terug, gewicht eraf.",
    progressie=(
        "Bij band: zwaardere band als je 15+ schone reps haalt. "
        "Bij dumbbell: 0,5-1 kg per 2 weken erbij. Doel binnen 6 weken: 4-5 kg."
    ),
)

exercise_block(
    nummer="K2",
    naam="Single-leg deadlift met dumbbell",
    materiaal="1 dumbbell (start 6-10 kg)",
    sets_reps="3 × 8 per been",
    tempo="3 sec naar voren kantelen, 2 sec terug",
    uitvoering=(
        "Sta op één been, dumbbell in de tegenoverliggende hand. Kantel vanuit de heup "
        "naar voren terwijl het achterste been gestrekt naar achter beweegt. "
        "Lichaam vormt één rechte lijn van kruin tot achterste hiel. "
        "Bekken horizontaal houden — laat de niet-belaste kant niet zakken."
    ),
    focus=(
        "Knie boven middelste teen — NIET naar binnen laten vallen. "
        "Bekken horizontaal — geen Trendelenburg-zakking."
    ),
    progressie="Per 2 weken 2 kg erbij. Of zwaarder: dumbbell in beide handen.",
)

exercise_block(
    nummer="K3",
    naam="Copenhagen plank",
    materiaal="Bank, stoel of bed",
    sets_reps="3 × 20-30 sec per kant",
    tempo="Statisch houden",
    uitvoering=(
        "Zijplank op de onderarm. Bovenste been op een bank/stoel (knie of enkel). "
        "Til je bekken op zodat je lichaam recht is. Onderste been hangt los — niet steunen."
    ),
    focus="Bekken hoog en horizontaal. Adembeweging blijft rustig.",
    progressie="Bouw op naar 45 sec; daarna onderste been aanlift voor 6-8 reps.",
)

exercise_block(
    nummer="K4",
    naam="Step-down voor spiegel",
    materiaal="Stap of trap van 20 cm hoogte. Eventueel dumbbells in de handen.",
    sets_reps="3 × 8 per been",
    tempo="3 sec omlaag, 1 sec onder, 2 sec omhoog",
    uitvoering=(
        "Sta op de stap met één been; laat het andere been langzaam zakken tot de hiel "
        "de grond raakt (niet erop steunen). Druk omhoog op het standbeen."
    ),
    focus=(
        "Spiegel ervoor — kijk live mee. Knie blijft BOVEN de middelste teen, valt niet "
        "naar binnen. Bekken blijft horizontaal. Als je dat niet kunt: tempo terug of "
        "stap-hoogte verlagen."
    ),
    progressie="Hoger opstapje (25-30 cm), of dumbbells in beide handen voor extra load.",
)

# ── WAT ERUIT MAG ─────────────────────────────────────────────────────────────
H1("Wat tijdelijk lichter mag")

P(
    "Zware split squats: tijdelijk terug naar ~70% van je huidige gewicht, met focus op "
    "kniespoor en horizontaal bekken. Zwaar trainen op een zwak fundament bouwt de zwakte "
    "erin. Eerst 4-6 weken medius-werk leveren, dan mag het gewicht in de split squat "
    "weer omhoog — én dan zal hij ook beter voelen."
)

# ── PROGRESSIE-MARKERS ───────────────────────────────────────────────────────
H1("Hoe je weet of het werkt")

P("Twee zelf-tests, elke twee weken:", bold=True)
P("")

doc.add_paragraph(
    "Single-leg squat tot 60° — film vanaf de voorkant. Zakt de knie naar binnen of "
    "blijft hij boven de middelste teen?",
    style="List Bullet",
)
doc.add_paragraph(
    "Trendelenburg-test — sta op één been voor de spiegel. Zakt de heup aan de "
    "niet-belaste kant naar beneden? Dat hoort niet.",
    style="List Bullet",
)
doc.add_paragraph(
    "Knie-twinge tijdens rustige Z2-runs: verdwijnt zonder dat je iets aan het lopen "
    "verandert. Dat is de echte marker.",
    style="List Bullet",
)

P("")
P("Streefbeeld na 6 weken:", bold=True)
P(
    "  • K1 (side-lying hip abduction) op 4-5 kg dumbbell of zwaarste band, 3×12 schoon."
)
P("  • K2 (single-leg deadlift) op 12-16 kg, bekken horizontaal door alle reps.")
P("  • Knie-twinge bij rustige loopjes verdwenen.")
P("  • Single-leg squat met knie keurig boven middelste teen, zonder bekkenkanteling.")

# ── REGELS ───────────────────────────────────────────────────────────────────
H1("Drie regels van Louis")

doc.add_paragraph("Tempo > gewicht. Vorm > volume.", style="List Number")
doc.add_paragraph(
    "Voel je TFL (zijkant bovenbeen) of onderrug overnemen? Stop. Beter 8 schone reps "
    "dan 15 vuile.",
    style="List Number",
)
doc.add_paragraph(
    "Drie keer per week aanraken weegt zwaarder dan één keer zwaar. Consistentie wint.",
    style="List Number",
)

doc.add_paragraph()
quote("\"De medius vuurt niet harder als je hem vermoeid traint. Hij vuurt harder als je hem schoon traint.\"")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("— Louis Delahaije")
run.italic = True

# ── OPSLAAN ──────────────────────────────────────────────────────────────────
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"Document opgeslagen: {OUT}")
